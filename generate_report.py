from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Styles helper ──
def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(0xD6, 0x28, 0x39)
        run.font.size = Pt(18)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        run.font.size = Pt(14)
    else:
        run.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)
        run.font.size = Pt(12)
    return p

def para(text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text).font.size = Pt(11)
    return p

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'D62839')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return t

# ═══════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("VeriCleri")
r.bold = True
r.font.size = Pt(36)
r.font.color.rgb = RGBColor(0xD6, 0x28, 0x39)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("AI-Powered Fake News Detection System")
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Final Year Project Documentation\nStudent: Hamza\nEmail: hamza02079@gmail.com\nLocation: Bahawalpur, Pakistan\n2026").font.size = Pt(12)

doc.add_page_break()

# ═══════════════════════════════════════
# 1. PROJECT OVERVIEW
# ═══════════════════════════════════════
heading("1. Project Overview")
para("VeriCleri is a web-based AI platform that detects whether a news article, headline, or WhatsApp message is REAL, FAKE, or UNCERTAIN. It uses Machine Learning trained on thousands of news samples — including Pakistani news data from Dawn, Geo, ARY News, and Express Tribune.")

doc.add_paragraph()
heading("Problem Statement", 2)
para("Fake news spreads rapidly on social media and WhatsApp, especially in Pakistan. People cannot easily verify news before sharing it. VeriCleri solves this by providing an instant, AI-powered fact-checking tool accessible from any browser — for free.")

heading("Objectives", 2)
for obj in [
    "Detect fake news using Machine Learning with high accuracy",
    "Support Pakistani news sources and local misinformation patterns",
    "Provide a simple, beautiful web interface for any user",
    "Include an AI chatbot (Veri) for interactive fact-checking",
    "Auto-retrain the model with fresh news data every 6 hours",
]:
    bullet(obj)

doc.add_page_break()

# ═══════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════
heading("2. System Architecture")
para("VeriCleri follows a 3-layer architecture:")

for item in [
    ("Frontend (Browser)", "HTML, CSS, JavaScript — what the user sees and interacts with"),
    ("Backend (Server)",   "Python + Flask — processes requests, runs ML model, returns results"),
    ("Data Layer",         "CSV datasets + JSON files — stores training data and statistics"),
]:
    p = doc.add_paragraph()
    p.add_run(f"{item[0]}: ").bold = True
    p.add_run(item[1])

doc.add_paragraph()
heading("Data Flow", 2)
steps = [
    "User pastes news text or URL in the browser",
    "Browser sends POST request to /analyze endpoint",
    "Flask backend receives the request",
    "If URL: scrapes article text using BeautifulSoup",
    "Rule-based scorer checks for fake/real language patterns",
    "ML model runs TF-IDF + Logistic Regression prediction",
    "Both scores combined → REAL / FAKE / UNCERTAIN verdict",
    "Result returned as JSON and displayed on result page",
]
for i, s in enumerate(steps, 1):
    bullet(f"Step {i}: {s}")

doc.add_page_break()

# ═══════════════════════════════════════
# 3. FRONTEND
# ═══════════════════════════════════════
heading("3. Frontend")
heading("Technologies Used", 2)
table(
    ["Technology", "Purpose"],
    [
        ["HTML5",           "Page structure and semantic markup"],
        ["CSS3 (Vanilla)",  "All styling, animations, dark mode"],
        ["JavaScript",      "Interactivity, API calls, animations"],
        ["Jinja2",          "Flask template engine (puts Python data into HTML)"],
        ["Font Awesome 6",  "Icons throughout the interface"],
        ["Google Fonts",    "Typography — Inter, Playfair Display"],
    ]
)

heading("Pages", 2)
table(
    ["Page", "File", "Purpose"],
    [
        ["Home",      "index.html",     "Landing page with hero section, stats, features"],
        ["Detect",    "detect.html",    "Main analysis tool — Text / URL / WhatsApp tabs"],
        ["Result",    "result.html",    "Shows ML verdict, confidence score, indicators"],
        ["Dashboard", "dashboard.html", "Live news feed from Pakistani & global RSS"],
        ["Base",      "base.html",      "Shared header, footer, Veri chatbot widget"],
    ]
)

heading("Key UI Features", 2)
for f in [
    "Dark Mode — toggled by button, saved in localStorage",
    "Responsive Design — works on mobile, tablet, and desktop",
    "Hero Animations — word-cycler, radar rings, scan beam, parallax tilt",
    "Scroll Reveal — IntersectionObserver animates elements as they appear",
    "Veri Chatbot — floating AI assistant (bottom-right corner)",
    "Confidence Count-up — number animates from 0 to final value on result page",
    "Progress Bars — animated fill bars showing fake/real probability",
    "Scanner Line — loading animation during analysis",
]:
    bullet(f)

doc.add_page_break()

# ═══════════════════════════════════════
# 4. BACKEND
# ═══════════════════════════════════════
heading("4. Backend")
heading("Technologies Used", 2)
table(
    ["Library / Framework", "Purpose"],
    [
        ["Flask",           "Web framework — handles HTTP routes and responses"],
        ["Python 3",        "Core programming language"],
        ["joblib",          "Save and load the trained ML model"],
        ["pandas",          "Load and manipulate CSV training data"],
        ["scikit-learn",    "Machine Learning (TF-IDF, Logistic Regression)"],
        ["BeautifulSoup4",  "Scrape article text from news URLs"],
        ["requests",        "Fetch web pages for URL analysis"],
        ["APScheduler",     "Background scheduler for auto-retraining"],
        ["flask-limiter",   "Rate limiting to prevent API abuse"],
        ["python-dotenv",   "Load secret keys from .env file"],
    ]
)

heading("API Endpoints", 2)
table(
    ["Route", "Method", "Purpose"],
    [
        ["/",                "GET",  "Home page"],
        ["/detect",          "GET",  "Detection tool page"],
        ["/analyze",         "POST", "Main ML analysis — accepts text or URL"],
        ["/result",          "GET",  "Result page (reads verdict from session)"],
        ["/dashboard",       "GET",  "Live news dashboard page"],
        ["/api/stats",       "GET",  "Returns site statistics as JSON"],
        ["/api/feedback",    "POST", "Stores user feedback on results"],
        ["/api/live-news",   "GET",  "Fetches RSS news for dashboard"],
        ["/api/retrain",     "POST", "Admin: triggers model retraining"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 5. MACHINE LEARNING MODEL
# ═══════════════════════════════════════
heading("5. Machine Learning Model")

heading("Algorithm", 2)
para("VeriCleri uses an Ensemble of two Logistic Regression pipelines:")
bullet("Word-level TF-IDF Pipeline — captures word patterns and phrases (1 to 3-grams)")
bullet("Character-level TF-IDF Pipeline — captures typos, punctuation abuse, writing style (3 to 5-grams)")
para("Both predictions are averaged to produce the final fake probability.")

heading("What is TF-IDF?", 2)
para("Term Frequency–Inverse Document Frequency is a mathematical formula that converts text into numbers. It gives high weight to words that appear often in one document but rarely across all documents. This helps the model identify unique 'fake news language' like 'shocking', 'secret', 'share before they delete'.")

heading("What is Logistic Regression?", 2)
para("A classification algorithm that outputs a probability between 0 and 1. In VeriCleri: 0 means REAL and 1 means FAKE. If the fake probability meets or exceeds the calibrated threshold, the verdict is FAKE.")

heading("Decision Thresholds (Calibrated)", 2)
table(
    ["Fake Probability", "Verdict"],
    [
        ["≥ 0.63 (63%)",       "FAKE"],
        ["≤ 0.46 (46%)",       "REAL"],
        ["Between 46% – 63%",  "UNCERTAIN"],
    ]
)

heading("Model Performance", 2)
table(
    ["Metric", "Score"],
    [
        ["Overall Accuracy",   "90.16%"],
        ["Decided Accuracy",   "96.32%  (when model is confident)"],
        ["F1 Score",           "90.17%"],
        ["ROC-AUC",            "97.41%"],
        ["Training Samples",   "4,145"],
        ["Test Samples",       "1,382"],
        ["Total Dataset Size", "5,527 samples"],
    ]
)

heading("Rule-Based Layer (on top of ML)", 2)
para("Fake Signals (+score):", bold=True)
for s in ["Sensational words: 'shocking', 'breaking', 'secret', 'cover-up'",
          "Misinformation phrases: 'vaccines cause autism', '5G causes COVID'",
          "Exclamation marks, ALL CAPS words",
          "Unsubstantiated death/war claims without citations",
          "'Share before they delete', 'urgent warning'"]:
    bullet(s)

para("Real Signals (+score):", bold=True)
for s in ["Citation phrases: 'according to', 'study shows', 'official report'",
          "Trusted sources: 'Reuters', 'BBC', 'Dawn', 'Geo', 'AP News'",
          "Numeric data, percentage figures, quoted sources",
          "Proper article length (80–1400 words)"]:
    bullet(s)

doc.add_page_break()

# ═══════════════════════════════════════
# 6. DATASETS
# ═══════════════════════════════════════
heading("6. Datasets")
table(
    ["Dataset", "Type", "Records"],
    [
        ["George McIntire (Kaggle)",     "Fake + Real articles",   "~44,000"],
        ["Lutz Fake-or-Real corpus",     "Fake news",              "~6,000"],
        ["Comprehensive fake (custom)",  "Curated fake samples",   "106"],
        ["Comprehensive real (custom)",  "Curated real samples",   "99"],
        ["Pakistan real news (custom)",  "Pakistani real news",    "120"],
        ["Pakistan fake news (custom)",  "Pakistani hoaxes",       "120"],
        ["Pakistan live RSS (scraped)",  "Live scraped headlines",  "282"],
        ["Latest reliable news (auto)",  "Auto-refreshed RSS",     "~250"],
    ]
)

heading("Pakistani News Sources", 2)
for s in ["Dawn (dawn.com)", "Geo News (geo.tv)", "ARY News (arynews.tv)",
          "The News (thenews.com.pk)", "Express Tribune (tribune.com.pk)",
          "Samaa TV (samaa.tv)", "Business Recorder (brecorder.com)"]:
    bullet(s)

heading("Pakistani Fake News Patterns Covered", 2)
for s in ["WhatsApp viral hoaxes ('Share before they delete')",
          "Political rumors (coup, arrest, secret deals)",
          "Financial scams (fake Ehsaas, BISP, NADRA schemes)",
          "Military/security panic (fake missile attacks, nuclear theft)",
          "Religious misinformation (Mecca-related fake news)"]:
    bullet(s)

doc.add_page_break()

# ═══════════════════════════════════════
# 7. VERI CHATBOT
# ═══════════════════════════════════════
heading("7. Veri — AI Chatbot")
para("Veri is a client-side chatbot widget (bottom-right corner). It does NOT use any external AI API like ChatGPT — everything runs locally.")

heading("How Veri Works", 2)
for s in [
    "User types a message in the chat window",
    "If it is a URL → calls /analyze with the URL → shows verdict",
    "If it is long text (12+ words) → calls /analyze → shows verdict in chat",
    "If it is a short question → keyword intent matching → returns pre-written answer",
]:
    bullet(s)

heading("Intents Veri Understands", 2)
table(
    ["User Says", "Veri Responds With"],
    [
        ["How does it work?",        "Explains TF-IDF + Logistic Regression pipeline"],
        ["How accurate?",            "Gives accuracy stats (90%, 96%, 97% AUC)"],
        ["WhatsApp message check",   "Guides to WhatsApp tab on Detection Desk"],
        ["Pakistani news / Salam",   "Explains Pakistani training data"],
        ["URL / article",            "Guides to URL tab on Detection Desk"],
        ["Hello / Hi / Assalam",     "Greets user in English or Urdu"],
    ]
)

heading("Why No ChatGPT/Gemini?", 2)
para("To keep VeriCleri free, private, and offline-capable — no API costs, no data sent to third parties, and FAQ questions work even without internet.")

doc.add_page_break()

# ═══════════════════════════════════════
# 8. AUTO-RETRAINING
# ═══════════════════════════════════════
heading("8. Auto-Retraining System")
para("The model retrains automatically every 6 hours using APScheduler:")
for s in [
    "Fetch latest headlines from 13 global RSS feeds (Reuters, BBC, CNN, NYT, AP, Guardian, Al Jazeera...)",
    "Also fetches from 7 Pakistani RSS feeds (Dawn, Geo, ARY, Tribune, Samaa...)",
    "Save new headlines as REAL news training samples",
    "Retrain the ML model with the updated combined dataset",
    "Save the new model to disk — old model replaced automatically",
]:
    bullet(s)

para("This means VeriCleri stays current — new language patterns from today's news are learned automatically without any manual work.")

doc.add_page_break()

# ═══════════════════════════════════════
# 9. SECURITY
# ═══════════════════════════════════════
heading("9. Security Features")
table(
    ["Feature", "How It Works"],
    [
        ["Rate Limiting",      "flask-limiter prevents API spam/abuse"],
        ["Input Validation",   "Text length capped at 15,000 characters"],
        ["Admin API Key",      "Retraining endpoints require X-Admin-Key header"],
        ["Secret Key",         "Loaded from .env file — never hardcoded in code"],
        ["XSS Prevention",     "Jinja2 auto-escapes all template variables"],
        ["No SQL = No SQLi",   "CSV-based storage — no SQL injection risk"],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════
# 10. PROJECT STRUCTURE
# ═══════════════════════════════════════
heading("10. Project File Structure")
code_lines = [
    "fakenews website/",
    "├── app.py                    ← Main Flask application (backend)",
    "├── requirements.txt          ← Python dependencies",
    "├── .env                      ← Secret keys (not on GitHub)",
    "│",
    "├── templates/                ← HTML pages (Jinja2)",
    "│   ├── base.html             ← Shared layout + Veri chatbot",
    "│   ├── index.html            ← Home page",
    "│   ├── detect.html           ← Detection tool",
    "│   ├── result.html           ← Analysis results",
    "│   └── dashboard.html        ← Live news feed",
    "│",
    "├── static/",
    "│   ├── css/style.css         ← All styles + animations",
    "│   ├── js/main.js            ← Core JS (dark mode, stats)",
    "│   └── js/chatbot.js         ← Veri chatbot logic",
    "│",
    "├── data/",
    "│   ├── sample_data.csv       ← Main training dataset",
    "│   ├── training/             ← Additional training CSVs",
    "│   └── raw_datasets/         ← External datasets for ingestion",
    "│",
    "├── ml_model/",
    "│   └── sample_models/",
    "│       └── *.joblib          ← Saved trained model",
    "│",
    "└── utils/",
    "    ├── dataset_ingest.py     ← CSV normalization for training",
    "    ├── pakistan_news_scraper.py ← Pakistani RSS scraper",
    "    ├── validators.py         ← Input validation",
    "    └── run_ingest_retrain.py ← Manual retrain trigger",
]
p = doc.add_paragraph()
p.style.font.name = 'Courier New'
for line in code_lines:
    doc.add_paragraph(line).runs[0].font.name = 'Courier New'

doc.add_page_break()

# ═══════════════════════════════════════
# 11. VIVA Q&A
# ═══════════════════════════════════════
heading("11. Viva / Exam Q&A Preparation")

qas = [
    ("Q: What problem does VeriCleri solve?",
     "Fake news spreads rapidly on social media and WhatsApp in Pakistan. People cannot verify news before sharing. VeriCleri provides an instant AI fact-checking tool accessible from any browser, for free."),

    ("Q: Why Machine Learning and not manual rules?",
     "Rule-based systems are brittle — you cannot manually write rules for every type of fake news. ML learns patterns automatically from thousands of examples and generalizes to new, unseen news it was never trained on."),

    ("Q: Why Logistic Regression and not Deep Learning?",
     "Logistic Regression with TF-IDF performs excellently on text classification. It is fast (under 2 seconds), interpretable, and achieves 97% ROC-AUC. Deep learning would require a GPU and much more data with no significant benefit here."),

    ("Q: What is TF-IDF?",
     "Term Frequency–Inverse Document Frequency converts text into numbers. Words appearing often in fake news (like 'shocking', 'secret') get high weights, helping the model distinguish fake from real."),

    ("Q: What is ROC-AUC?",
     "Receiver Operating Characteristic — Area Under Curve measures how well the model separates the two classes (real vs fake). A score of 97.41% means excellent discrimination between real and fake news."),

    ("Q: What is the UNCERTAIN verdict?",
     "When the fake probability falls between 46%–63%, the model lacks enough confidence. Instead of making a wrong confident prediction, VeriCleri honestly says 'uncertain — please verify manually'."),

    ("Q: How did you handle Pakistani news?",
     "We created 240 hand-curated Pakistani samples (120 real, 120 fake) covering Pakistani institutions, politicians, cities, and common hoax patterns. We also built an RSS scraper that fetches live headlines from Dawn, Geo, ARY, Express Tribune, and Samaa TV."),

    ("Q: How does URL analysis work?",
     "When a user submits a URL, the backend uses the requests library to fetch the page HTML and BeautifulSoup to extract the main article text by stripping navigation and ads. This text is then passed to the same ML pipeline."),

    ("Q: What are the limitations?",
     "1) Works best with English (Urdu/Roman Urdu support is limited). 2) Very short texts under 12 words are harder to classify. 3) Satire can sometimes be misclassified. 4) Brand-new misinformation types not in training data may be missed."),

    ("Q: What is Flask?",
     "Flask is a lightweight Python web framework. It handles HTTP requests, routes URLs to Python functions, renders HTML templates, and sends responses back to the browser. It connects the frontend to the ML model."),

    ("Q: Why not use a database?",
     "For this project, CSV files and JSON are sufficient for storing training data, feedback, and stats. This keeps the project simple and portable — no database server needed. In production, PostgreSQL or MongoDB would be used."),

    ("Q: What is the chatbot Veri?",
     "Veri is a client-side chatbot. It uses keyword-based intent matching for FAQ answers and calls the site's own /analyze API for news analysis. No external AI service is used — free, private, and offline-capable for FAQ questions."),

    ("Q: How does dark mode work?",
     "A CSS class 'dark-mode' is toggled on the body element using JavaScript. CSS variables define both light and dark theme colors. The user's preference is saved in localStorage so it persists across page visits."),

    ("Q: How is the model saved?",
     "Using joblib — a Python library optimized for saving large numpy arrays and scikit-learn objects. The trained model (word pipeline + char pipeline + metrics) is serialized to a .joblib binary file and loaded back into memory when the server starts."),

    ("Q: What is GridSearchCV?",
     "An automated hyperparameter tuning technique. It tries multiple values of the regularization parameter C (0.5, 1.0, 2.0, 4.0) with cross-validation and picks the combination that gives the best F1 score on validation data."),
]

for q, a in qas:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    p.runs[0].font.color.rgb = RGBColor(0xD6, 0x28, 0x39)
    p.runs[0].font.size = Pt(11)
    ans = doc.add_paragraph()
    ans.add_run(a).font.size = Pt(11)
    doc.add_paragraph()

doc.add_page_break()

# ═══════════════════════════════════════
# 12. HOW TO RUN
# ═══════════════════════════════════════
heading("12. How to Run the Project")
steps = [
    ("Install dependencies",  "pip install -r requirements.txt"),
    ("Create .env file",      'echo SECRET_KEY=any_random_string > .env'),
    ("Start the server",      "python app.py"),
    ("Open in browser",       "http://localhost:5000"),
]
for title, cmd in steps:
    p = doc.add_paragraph()
    p.add_run(f"{title}: ").bold = True
    r = p.add_run(cmd)
    r.font.name = 'Courier New'
    r.font.color.rgb = RGBColor(0xD6, 0x28, 0x39)

# ═══════════════════════════════════════
# SAVE
# ═══════════════════════════════════════
path = r"C:\Users\hamza\Desktop\VeriCleri_FYP_Documentation.docx"
doc.save(path)
print(f"[OK] Word file saved: {path}")
